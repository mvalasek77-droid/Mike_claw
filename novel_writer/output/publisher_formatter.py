"""
Publisher-Ready Formatter
Converts the humanized manuscript into industry-standard submission formats:
- .docx  — standard manuscript format (Times New Roman 12pt, double-spaced, 1" margins)
- .txt   — clean plain text
- .md    — Markdown for digital review

Industry manuscript standards per Chicago Manual of Style + Writer's Digest:
- Times New Roman 12pt
- Double line spacing
- 1-inch margins all around
- Header: Author Last Name / Short Title / Page Number (top right)
- First page: Title, Author, word count (centered)
- Scene breaks: ###
- Chapter titles: centered, no special formatting
- ~250 words per page (standard manuscript estimate)
"""

from __future__ import annotations
import re
import json
import math
from pathlib import Path
from typing import Any
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import (
    MANUSCRIPT_FONT, MANUSCRIPT_FONT_SIZE, MANUSCRIPT_LINE_SPACE,
    OUTPUT_FORMATS,
)

OUTPUT_DIR = Path(__file__).parent


def count_words(text: str) -> int:
    return len(text.split())


def estimate_pages(word_count: int, words_per_page: int = 250) -> int:
    return math.ceil(word_count / words_per_page)


def clean_chapter_text(text: str) -> str:
    """Normalize whitespace, fix common encoding artifacts."""
    # Normalize quotes
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    # Normalize dashes
    text = text.replace("\u2014", "—").replace("\u2013", "–")
    # Normalize ellipses
    text = text.replace("\u2026", "...")
    # Remove excessive blank lines (max 2 consecutive)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Plain text formatter ───────────────────────────────────────────────────────

def format_txt(
    chapters: list[dict[str, Any]],
    novel_plan: dict[str, Any],
    output_path: Path,
) -> Path:
    """Write a clean plain text manuscript."""
    title = novel_plan.get("title", "UNTITLED")
    author = novel_plan.get("author", "Anonymous")
    total_words = sum(c.get("word_count", count_words(c["text"])) for c in chapters)
    est_pages = estimate_pages(total_words)

    lines = [
        title.upper(),
        "",
        f"By {author}",
        "",
        f"Approximately {total_words:,} words ({est_pages:,} pages)",
        "",
        "=" * 60,
        "",
    ]

    for i, ch in enumerate(chapters, 1):
        lines.append("")
        lines.append(ch.get("title", f"CHAPTER {i}").upper())
        lines.append("")
        lines.append(clean_chapter_text(ch["text"]))
        lines.append("")
        lines.append("# # #")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"  [format] TXT → {output_path}")
    return output_path


# ── Markdown formatter ─────────────────────────────────────────────────────────

def format_md(
    chapters: list[dict[str, Any]],
    novel_plan: dict[str, Any],
    output_path: Path,
) -> Path:
    """Write Markdown manuscript with metadata front matter."""
    title = novel_plan.get("title", "Untitled")
    author = novel_plan.get("author", "Anonymous")
    total_words = sum(c.get("word_count", count_words(c["text"])) for c in chapters)

    lines = [
        "---",
        f'title: "{title}"',
        f'author: "{author}"',
        f'genre: "{novel_plan.get("genre", "Literary Fiction")}"',
        f'words: {total_words}',
        f'logline: "{novel_plan.get("logline", "")}"',
        "---",
        "",
        f"# {title}",
        "",
        f"*{novel_plan.get('back_cover_copy', '')}*",
        "",
        "---",
        "",
    ]

    for i, ch in enumerate(chapters, 1):
        lines.append(f"## {ch.get('title', f'Chapter {i}')}")
        lines.append("")
        lines.append(clean_chapter_text(ch["text"]))
        lines.append("")
        lines.append("---")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"  [format] MD  → {output_path}")
    return output_path


# ── DOCX formatter ─────────────────────────────────────────────────────────────

def format_docx(
    chapters: list[dict[str, Any]],
    novel_plan: dict[str, Any],
    output_path: Path,
) -> Path:
    """
    Write a publisher-standard .docx manuscript.
    Requires python-docx. Falls back gracefully if not installed.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        print("  [format] python-docx not installed. Skipping .docx output.")
        print("           Install with: pip install python-docx")
        return None

    title = novel_plan.get("title", "Untitled")
    author = novel_plan.get("author", "Anonymous")
    author_last = author.split()[-1] if author != "Anonymous" else "Author"
    title_short = " ".join(title.split()[:3])
    total_words = sum(c.get("word_count", count_words(c["text"])) for c in chapters)

    doc = Document()

    # ── Page setup (1-inch margins) ────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Default style (Times New Roman 12pt, double-spaced) ───────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = MANUSCRIPT_FONT
    font.size = Pt(MANUSCRIPT_FONT_SIZE)
    pf = style.paragraph_format
    pf.line_spacing = MANUSCRIPT_LINE_SPACE * Pt(MANUSCRIPT_FONT_SIZE)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    def add_paragraph(text: str = "", style_name: str = "Normal", alignment=None) -> Any:
        p = doc.add_paragraph(text, style=style_name)
        if alignment is not None:
            p.alignment = alignment
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = MANUSCRIPT_FONT
        run.font.size = Pt(MANUSCRIPT_FONT_SIZE)
        return p

    # ── Title page ─────────────────────────────────────────────────────────────
    # Contact info block (top left)
    for _ in range(10):
        doc.add_paragraph("")

    p_title = add_paragraph(title.upper(), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p_title.runs[0].bold = True
    p_title.runs[0].font.size = Pt(14)

    add_paragraph("", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(f"by", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(author, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(8):
        doc.add_paragraph("")

    add_paragraph(
        f"Approximately {total_words:,} words",
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    # ── Header (Author / Title / Page#) ───────────────────────────────────────
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run(f"{author_last} / {title_short} / ")
    run.font.name = MANUSCRIPT_FONT
    run.font.size = Pt(MANUSCRIPT_FONT_SIZE)
    # Add page number field
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = header_para.add_run()
    run2._r.append(fldChar)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run2.font.name = MANUSCRIPT_FONT
    run2.font.size = Pt(MANUSCRIPT_FONT_SIZE)

    # ── Chapters ──────────────────────────────────────────────────────────────
    for i, ch in enumerate(chapters, 1):
        doc.add_page_break()

        # Chapter title (centered, ~1/3 down page)
        for _ in range(10):
            doc.add_paragraph("")
        ch_title_para = add_paragraph(
            ch.get("title", f"Chapter {i}").upper(),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        ch_title_para.runs[0].bold = False
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Chapter body
        text = clean_chapter_text(ch["text"])
        paragraphs = text.split("\n\n")

        for j, para_text in enumerate(paragraphs):
            para_text = para_text.strip()
            if not para_text:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = MANUSCRIPT_LINE_SPACE * Pt(MANUSCRIPT_FONT_SIZE)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # First paragraph of chapter: no indent
            # Subsequent paragraphs: 0.5-inch first-line indent
            if j > 0:
                p.paragraph_format.first_line_indent = Inches(0.5)

            run = p.add_run(para_text)
            run.font.name = MANUSCRIPT_FONT
            run.font.size = Pt(MANUSCRIPT_FONT_SIZE)

    doc.save(str(output_path))
    print(f"  [format] DOCX → {output_path}")
    return output_path


# ── Manuscript statistics ─────────────────────────────────────────────────────

def generate_manuscript_report(
    chapters: list[dict[str, Any]],
    novel_plan: dict[str, Any],
    output_path: Path,
) -> dict[str, Any]:
    """Generate a submission-ready synopsis + statistics sheet."""
    total_words = sum(c.get("word_count", count_words(c["text"])) for c in chapters)

    stats = {
        "title": novel_plan.get("title"),
        "genre": novel_plan.get("genre"),
        "setting": novel_plan.get("setting"),
        "total_word_count": total_words,
        "estimated_pages": estimate_pages(total_words),
        "chapter_count": len(chapters),
        "logline": novel_plan.get("logline"),
        "comparable_titles": novel_plan.get("comparable_titles", []),
        "back_cover_copy": novel_plan.get("back_cover_copy"),
        "chapter_breakdown": [
            {
                "chapter": c["chapter"],
                "title": c.get("title", f"Chapter {c['chapter']}"),
                "word_count": c.get("word_count", count_words(c["text"])),
                "beat": c.get("beat", ""),
            }
            for c in chapters
        ],
        "ai_humanization": {
            "avg_score_pre": sum(c.get("ai_score_pre", 0) for c in chapters) / len(chapters) if chapters else 0,
            "avg_score_post": sum(c.get("ai_score_post", 0) for c in chapters) / len(chapters) if chapters else 0,
        },
    }

    with open(output_path, "w") as f:
        json.dump(stats, f, indent=2)

    # Also write a text synopsis
    synopsis_path = output_path.with_suffix(".txt")
    with open(synopsis_path, "w") as f:
        f.write(f"MANUSCRIPT SUBMISSION PACKAGE\n{'='*50}\n\n")
        f.write(f"TITLE: {stats['title']}\n")
        f.write(f"GENRE: {stats['genre']}\n")
        f.write(f"WORD COUNT: {stats['total_word_count']:,}\n")
        f.write(f"ESTIMATED PAGES: {stats['estimated_pages']:,}\n\n")
        f.write(f"LOGLINE:\n{stats['logline']}\n\n")
        f.write(f"COMPARABLE TITLES:\n")
        for ct in stats['comparable_titles']:
            f.write(f"  • {ct}\n")
        f.write(f"\nBACK COVER COPY:\n{stats['back_cover_copy']}\n\n")
        f.write(f"CHAPTER BREAKDOWN:\n")
        for ch in stats['chapter_breakdown']:
            f.write(f"  Ch {ch['chapter']:02d}: {ch['title']} ({ch['word_count']:,}w) — {ch['beat']}\n")

    print(f"  [format] Report → {output_path}")
    print(f"  [format] Synopsis → {synopsis_path}")
    return stats


# ── Master formatter ──────────────────────────────────────────────────────────

class PublisherFormatter:
    def __init__(self, output_dir: Path | None = None):
        self.output_dir = output_dir or OUTPUT_DIR

    def format_all(
        self,
        chapters: list[dict[str, Any]],
        novel_plan: dict[str, Any],
    ) -> dict[str, str]:
        """Write all output formats. Returns dict of {format: path}."""
        title_slug = re.sub(r"[^a-z0-9]+", "_", novel_plan.get("title", "untitled").lower())
        outputs = {}

        print(f"\n  [format] Generating publisher-ready outputs...")

        # TXT
        txt_path = self.output_dir / f"{title_slug}.txt"
        format_txt(chapters, novel_plan, txt_path)
        outputs["txt"] = str(txt_path)

        # Markdown
        md_path = self.output_dir / f"{title_slug}.md"
        format_md(chapters, novel_plan, md_path)
        outputs["md"] = str(md_path)

        # DOCX
        docx_path = self.output_dir / f"{title_slug}.docx"
        result = format_docx(chapters, novel_plan, docx_path)
        if result:
            outputs["docx"] = str(docx_path)

        # Report
        report_path = self.output_dir / f"{title_slug}_report.json"
        stats = generate_manuscript_report(chapters, novel_plan, report_path)
        outputs["report"] = str(report_path)

        total_words = stats["total_word_count"]
        print(f"\n  [format] ✓ Complete!")
        print(f"  [format] {total_words:,} words | {stats['estimated_pages']:,} pages | {len(chapters)} chapters")

        return outputs


def format_manuscript(
    chapters: list[dict[str, Any]],
    novel_plan: dict[str, Any],
    output_dir: Path | None = None,
) -> dict[str, str]:
    """Convenience entry point called by the orchestrator."""
    formatter = PublisherFormatter(output_dir)
    return formatter.format_all(chapters, novel_plan)
